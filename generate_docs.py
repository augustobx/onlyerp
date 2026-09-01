import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", 4)}" w:space="{edge_data.get("space", 0)}" w:color="{edge_data.get("color", "auto")}"/>'
            tcBorders.append(parse_xml(tag))
        else:
            tag = f'<w:{edge} {nsdecls("w")} w:val="none"/>'
            tcBorders.append(parse_xml(tag))
    tcPr.append(tcBorders)

def add_callout(doc, text, title="NOTA IMPORTANTE", bg_color="F0F4F8", border_color="3B49DF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    set_cell_border(cell, 
                    left=dict(sz=24, val="single", color=border_color),
                    top=dict(sz=4, val="single", color="E2E8F0"),
                    bottom=dict(sz=4, val="single", color="E2E8F0"),
                    right=dict(sz=4, val="single", color="E2E8F0"))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run_title = p.add_run(f"📌 {title}: ")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor.from_string(border_color)
    
    run_text = p.add_run(text)
    run_text.font.name = "Segoe UI"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    doc.add_paragraph()

def add_header_banner(doc, tag, title, subtitle, bg_color="1E1B4B", tag_color="818CF8"):
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    header_cell.width = Inches(6.9)
    set_cell_background(header_cell, bg_color)
    set_cell_margins(header_cell, top=280, bottom=280, left=250, right=250)
    set_cell_border(header_cell)
    
    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(4)
    
    tag_run = hp.add_run(tag)
    tag_run.font.name = "Segoe UI"
    tag_run.font.size = Pt(9)
    tag_run.font.bold = True
    tag_run.font.color.rgb = RGBColor.from_string(tag_color)
    
    hp2 = header_cell.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp2.paragraph_format.space_after = Pt(6)
    
    title_run = hp2.add_run(title)
    title_run.font.name = "Segoe UI"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    hp3 = header_cell.add_paragraph()
    hp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp3.paragraph_format.space_after = Pt(0)
    
    sub_run = hp3.add_run(subtitle)
    sub_run.font.name = "Segoe UI"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)
    
    doc.add_paragraph()

def add_section_header(doc, number, title, subtitle=None, primary_color="4F46E5"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    num_run = p.add_run(f"{number}. ")
    num_run.font.name = "Segoe UI"
    num_run.font.size = Pt(14)
    num_run.font.bold = True
    num_run.font.color.rgb = RGBColor.from_string(primary_color)
    
    t_run = p.add_run(title)
    t_run.font.name = "Segoe UI"
    t_run.font.size = Pt(14)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    if subtitle:
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after = Pt(8)
        ps.paragraph_format.keep_with_next = True
        s_run = ps.add_run(subtitle)
        s_run.font.name = "Segoe UI"
        s_run.font.size = Pt(9.5)
        s_run.font.italic = True
        s_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def add_sub_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    t_run = p.add_run(title)
    t_run.font.name = "Segoe UI"
    t_run.font.size = Pt(11)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(bold_prefix + " ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def add_step(doc, num_str, bold_prefix, text, primary_color="4F46E5"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    num_run = p.add_run(f"Paso {num_str}: ")
    num_run.bold = True
    num_run.font.color.rgb = RGBColor.from_string(primary_color)
    
    r1 = p.add_run(bold_prefix + " ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def init_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return doc

# ==============================================================================
# MANUAL DEL SISTEMA ADMINISTRATIVO (COMPLETO CON GUÍA DETALLADA DE ESTADOS)
# ==============================================================================
def generar_manual_sistema():
    doc = init_doc()
    
    add_header_banner(
        doc,
        tag="MANUAL OPERATIVO • ADMINISTRACIÓN Y OFICINA",
        title="Guía de Uso del Sistema Web Central",
        subtitle="Caja Diaria • Ciclo de Vida de Pedidos • Cuentas Corrientes • Promociones y Combos",
        bg_color="0F172A",
        tag_color="38BDF8"
    )
    
    # --------------------------------------------------------------------------
    # 1. CAJA DIARIA
    # --------------------------------------------------------------------------
    add_section_header(doc, "1", "Apertura y Cierre de Caja Diaria", "Control estricto de turnos, arqueos físicos ciegos y conciliación financiera.")
    add_sub_header(doc, "Apertura de Turno")
    add_step(doc, "1", "Acceso al módulo:", "Ingresá a 'Caja' desde el menú lateral del sistema.")
    add_step(doc, "2", "Estado cerrado:", "Cuando el turno está cerrado, la facturación del punto de venta se encuentra bloqueada.")
    add_step(doc, "3", "Declarar Fondo Inicial:", "En el campo 'Declarar Efectivo Inicial ($)', ingresá el monto exacto de cambio/fondo en el cajón.")
    add_step(doc, "4", "Abrir turno:", "Presioná 'Abrir Turno Ahora'. El sistema registrará la fecha, hora exacta y el cajero activo.")

    add_sub_header(doc, "Métricas y Monitoreo en Tiempo Real")
    add_bullet(doc, "Efectivo en Cajón:", "Saldo inicial + cobranzas y ventas en efectivo - egresos manuales.")
    add_bullet(doc, "Transferencias:", "Total de cobros ingresados por cuentas bancarias o billeteras virtuales.")
    add_bullet(doc, "Tarjetas:", "Ventas y cobros registrados mediante terminales de débito o crédito.")
    add_bullet(doc, "Salidas / Gastos:", "Total de dinero en efectivo retirado del cajón durante la jornada.")

    add_sub_header(doc, "Registro de Retiros o Gastos")
    add_step(doc, "1", "Botón 'Retiro / Gasto':", "Hacé clic en el botón superior de la vista de caja.")
    add_step(doc, "2", "Completar motivo:", "Ingresá el monto y el concepto del egreso (ej: 'Pago de flete', 'Compra de librería').")
    add_step(doc, "3", "Guardar:", "Presioná 'Guardar Retiro'. El dinero se descontará de la caja y quedará auditado en el historial.")

    add_sub_header(doc, "Cierre de Turno y Arqueo Ciego")
    add_step(doc, "1", "Finalizar turno:", "Hacé clic en 'Cerrar Turno'.")
    add_step(doc, "2", "Declaración física:", "Contá el dinero real del cajón e ingresalo en 'Efectivo Físico Contado ($)'.")
    add_step(doc, "3", "Balance:", "Presioná 'Finalizar Turno'. El sistema calculará la Diferencia (+ / -) y la Ganancia neta generada.")
    
    add_callout(doc, "El sistema almacena el 'Historial de Turnos Anteriores' con fecha, hora, cajero, diferencias y ganancia, permitiendo auditar cada movimiento individualmente.", "CONTROL DE AUDITORÍA")

    # --------------------------------------------------------------------------
    # 2. CICLO DE VIDA DEL PEDIDO (DETALLE EXPLICADO Y TRANSICIONES)
    # --------------------------------------------------------------------------
    add_section_header(doc, "2", "Ciclo de Vida del Pedido: Flujo Operativo y Cambio de Estados", "Cómo sigue un pedido desde su ingreso hasta la entrega final y el impacto en cada etapa.")
    
    add_sub_header(doc, "Esquema General del Flujo")
    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.space_before = Pt(2)
    p_flow.paragraph_format.space_after = Pt(6)
    p_flow.paragraph_format.line_spacing = 1.15
    r_flow = p_flow.add_run("1. PENDIENTE  ──>  2. APROBADO  ──>  3. ARMADO (Despacho)  ──>  4. FACTURADO  ──>  5. ENTREGADO")
    r_flow.bold = True
    r_flow.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    # Table of states summary
    states_table = doc.add_table(rows=1, cols=3)
    states_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    states_table.autofit = False
    
    hdr_cells = states_table.rows[0].cells
    hdr_cells[0].width = Inches(1.5)
    hdr_cells[1].width = Inches(2.2)
    hdr_cells[2].width = Inches(2.8)
    
    headers = ["Estado", "Significado / Etapa", "Impacto en Stock y Sistema"]
    for i, h in enumerate(headers):
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    states_data = [
        ("PENDIENTE", "Pedido ingresado desde la app o cargado por oficina.", "Reserva preventiva inmediata de stock para evitar sobreventas."),
        ("APROBADO", "Revisión comercial y crediticia aprobada.", "Habilita la orden de preparación física en depósito."),
        ("ARMADO / LISTO ENTREGA", "Mercadería embalada con repartidor y fecha asignada.", "Visible en Despacho y en la app móvil del transportista."),
        ("FACTURADO", "Emisión de comprobante AFIP (A/B/C) o Comprobante X.", "Crea la Venta definitiva. Impacta en Cuenta Corriente o en Caja."),
        ("ENTREGADO", "Mercadería entregada al cliente con éxito.", "Finalización del circuito logístico."),
        ("NO ENTREGADO", "Incidencia (local cerrado, cliente ausente).", "Registro del motivo y posibilidad de reprogramar."),
        ("CANCELADO / RECHAZADO", "Pedido anulado por vendedor o administración.", "Libera y devuelve el stock preventivo al inventario disponible al instante.")
    ]

    for row_idx, (st, sig, imp) in enumerate(states_data):
        row_cells = states_table.add_row().cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([st, sig, imp]):
            cell = row_cells[i]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            set_cell_border(cell, bottom=dict(sz=2, val="single", color="E2E8F0"))
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.name = "Segoe UI"
            r.font.size = Pt(8.5)
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            else:
                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()

    # DETAILED STEP BY STEP TRANSITION EXPLANATION
    add_sub_header(doc, "Explicación Paso a Paso: Cómo debe Avanzar el Pedido")

    # ETAPA 1: PENDIENTE
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("Etapa 1: Recepción e Inspección (Estado: PENDIENTE)")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    add_bullet(doc, "Origen:", "El pedido entra automáticamente desde la PWA del vendedor en la calle o se genera en el sistema.")
    add_bullet(doc, "Reserva de Stock:", "El sistema descuenta preventivamente las unidades del depósito central. Ningún otro usuario podrá vender mercadería ya comprometida.")
    add_bullet(doc, "Acciones posibles en esta etapa:", "")
    add_step(doc, "A", "Editar Pedido:", "Hacé clic en 'Editar' para cambiar cantidades, añadir o quitar productos, aplicar descuentos por ítem o modificar notas internas.", "0284C7")
    add_step(doc, "B", "Recalcular Precios:", "Si los costos o listas de precios cambiaron recientemente, presioná 'Recalcular' en la cabecera para actualizar los montos de todos los pendientes.", "0284C7")
    add_step(doc, "C", "Rechazar:", "Si el pedido es inviable o no hay acuerdo con el cliente, hacé clic en 'Rechazar'. El sistema liberará y devolverá el stock al inventario de inmediato.", "0284C7")
    add_step(doc, "D", "Aprobar:", "Presioná 'Aprobar' para validar comercialmente la operación y pasar a preparación.", "0284C7")

    # ETAPA 2: APROBADO
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Etapa 2: Autorización Comercial y Preparación (Estado: APROBADO)")
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    add_bullet(doc, "Objetivo:", "El pedido cuenta con el visto bueno de administración (límite de crédito verificado, condiciones comerciales aceptadas) y el depósito recibe la orden de armar los bultos.")
    add_bullet(doc, "Transición siguiente:", "Hacé clic en el botón 'Listo p/ Entrega (Armar)' para transferir el pedido al área de logística y despacho.")

    # ETAPA 3: ARMADO / LISTO ENTREGA
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run("Etapa 3: Armado Físico y Asignación de Despacho (Estado: ARMADO / LISTO_ENTREGA)")
    r3.bold = True
    r3.font.size = Pt(10.5)
    r3.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    add_bullet(doc, "Módulo de Armados / Despacho (`/pedidos/armados`):", "Permite gestionar las órdenes listas para ser cargadas en los vehículos.")
    add_bullet(doc, "Asignación logística:", "Se define el Repartidor/Chofer responsable y la Fecha de Entrega exacta.")
    add_bullet(doc, "Sincronización móvil:", "El pedido aparece instantáneamente en la pestaña 'Repartos' del chofer o vendedor en su app de celular.")

    # ETAPA 4: FACTURADO
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(8)
    p4.paragraph_format.space_after = Pt(2)
    r4 = p4.add_run("Etapa 4: Emisión Fiscal o Comercial (Acción: FACTURAR)")
    r4.bold = True
    r4.font.size = Pt(10.5)
    r4.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    add_bullet(doc, "¿Cuándo facturar?:", "Podés facturar el pedido cuando está Aprobado, Armado o antes de que salga del depósito presionando el botón 'Facturar Pedido'.")
    add_bullet(doc, "Modal de Facturación inteligente:", "El sistema analiza la condición impositiva del cliente:")
    add_bullet(doc, "• Responsable Inscripto / Monotributo:", "Sugiere Factura A/B/C y se conecta con el Web Service de AFIP para obtener CAE oficial y fecha de vencimiento.")
    add_bullet(doc, "• Consumidor Final / Interno:", "Genera Comprobante X.")
    add_bullet(doc, "Impacto financiero automático:", "")
    add_bullet(doc, "• Si el pago es Cuenta Corriente:", "Carga el saldo a la ficha del cliente con vencimiento automático.")
    add_bullet(doc, "• Si el pago es Contado / Efectivo:", "Registra el ingreso en la Caja Diaria abierta de la sucursal.")
    add_bullet(doc, "Impresión:", "Habilita los botones de impresión de Ticket térmico de 80mm o Factura oficial en formato A4.")

    # ETAPA 5: ENTREGADO / NO ENTREGADO
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(8)
    p5.paragraph_format.space_after = Pt(2)
    r5 = p5.add_run("Etapa 5: Entrega en Destino e Incidencias (Estados: ENTREGADO / NO_ENTREGADO)")
    r5.bold = True
    r5.font.size = Pt(10.5)
    r5.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    add_bullet(doc, "Caso A - Entrega Exitosa (ENTREGADO):", "El repartidor desde su app (o el operador desde el sistema) presiona 'Marcar Entregado'. El pedido concluye su circuito logístico.")
    add_bullet(doc, "Caso B - Incidencia en Entrega (NO ENTREGADO):", "Si el cliente no estaba, el local estaba cerrado o se rechazó la entrega:")
    add_step(doc, "1", "Registrar No Entrega:", "Se presiona 'No Entregado' y se selecciona el motivo (Cliente ausente, Local cerrado, Rechazó mercadería, etc.).", "0284C7")
    add_step(doc, "2", "Reintentar Despacho:", "El pedido queda marcado en rojo con la incidencia visible. Desde la botonera se puede presionar 'Reintentar y Armar' para reprogramar la entrega al día siguiente sin perder la reserva de mercadería.", "0284C7")

    add_callout(doc, "Regla estricta de anulación: Si un pedido ya fue facturado, no puede cancelarse directamente desde Pedidos. Primero debe anularse el comprobante/factura desde el módulo de Historial de Ventas para mantener la consistencia fiscal y contable.", "SEGURIDAD FISCAL")

    # --------------------------------------------------------------------------
    # 3. CUENTAS CORRIENTES
    # --------------------------------------------------------------------------
    add_section_header(doc, "3", "Cobranza de Cuentas Corrientes", "Gestión de deudores, actualización automática por inflación y cobros.")
    add_sub_header(doc, "Listado de Saldos y Recordatorio por WhatsApp")
    add_step(doc, "1", "Panel General:", "Ingresá a 'Cuentas Corrientes' para ver el total adeudado en la calle y la cantidad de clientes deudores.")
    add_step(doc, "2", "Filtros:", "Buscá por nombre, DNI/CUIT o filtrá por 'Solo Vencidas' o 'Al Día'.")
    add_step(doc, "3", "Recordatorio WhatsApp:", "Hacé clic en el ícono de WhatsApp para enviar un mensaje predeterminado con el saldo deudor exacto.")

    add_sub_header(doc, "Ficha del Cliente y Recálculo por Inflación")
    add_step(doc, "1", "Abrir Ficha:", "Hacé clic en 'Ver Cuenta' sobre el cliente.")
    add_step(doc, "2", "Actualizar Inflación:", "Si la factura venció su plazo acordado, aparecerá el botón naranja 'Actualizar Inflación' para reajustar los precios a valores actuales.")
    add_step(doc, "3", "Registrar Cobro:", "Hacé clic en 'Cobrar', indicá el monto (parcial o total), medio de pago (Efectivo, Tarjeta o Transferencia) y descuento por pronto pago si correspondiera.")

    # --------------------------------------------------------------------------
    # 4. COMBOS
    # --------------------------------------------------------------------------
    add_section_header(doc, "4", "Creación de Promociones y Combos", "Armado de paquetes comerciales cerrados con descuentos o precios especiales.")
    add_step(doc, "1", "Acceso al módulo:", "Ingresá a 'Combos' (`/combos`) en el menú principal.")
    add_step(doc, "2", "Nuevo Combo:", "Presioná el botón '+ Nuevo Combo / Promo'.")
    add_step(doc, "3", "Datos comerciales:", "Indicá Nombre, Descripción, Precio Total del Combo ($) y % de Descuento promocional.")
    add_step(doc, "4", "Agregar productos:", "Buscá los artículos y definí la cantidad de unidades de cada producto que componen el paquete.")
    add_step(doc, "5", "Guardar:", "Al guardar, el combo se reflejará instantáneamente en la app de todos los vendedores.")

    output_path = r"c:\xampp\htdocs\tomassidok\Manual_de_Uso_Sistema_Administrativo.docx"
    doc.save(output_path)
    print(f"Manual del Sistema guardado en: {output_path}")

# ==============================================================================
# MANUAL DE LA APP DE VENDEDORES
# ==============================================================================
def generar_manual_app():
    doc = init_doc()
    
    add_header_banner(
        doc,
        tag="MANUAL OPERATIVO • VENDEDORES Y REPARTIDORES",
        title="Guía de Uso de la App Móvil (PWA)",
        subtitle="Toma de Pedidos en Calle • Catálogo • Combos • Descuentos • Repartos • Cobranzas",
        bg_color="312E81",
        tag_color="A5B4FC"
    )
    
    # 1. Introducción y Conectividad
    add_section_header(doc, "1", "Conectividad y Modo Offline", "Operación garantizada en la calle con o sin conexión a internet.")
    add_bullet(doc, "Indicador Verde (CONECTADO):", "Los pedidos y cobros se envían al instante a la base de datos central.")
    add_bullet(doc, "Indicador Rojo (MODO OFFLINE):", "Si estás sin señal o en zona rural, la app guarda los clientes y pedidos en la memoria del dispositivo y los sincroniza automáticamente al recuperar internet.")
    add_bullet(doc, "Botón Refrescar (Flechas circulares):", "Fuerza la sincronización inmediata de pedidos pendientes y actualiza stock.")

    # 2. Confección del Pedido
    add_section_header(doc, "2", "Paso a Paso: Toma de Pedidos", "Flujo de trabajo para armar y enviar una orden de venta.")
    
    add_sub_header(doc, "A. Cómo Elegir o Dar de Alta un Cliente")
    add_step(doc, "1", "Buscar Cliente:", "En la pestaña 'Nuevo', escribí el nombre o CUIT en la barra de búsqueda y tocalo en la lista.")
    add_step(doc, "2", "Crear Cliente Nuevo:", "Tocá '+ Nuevo Cliente', completá Nombre, CUIT/DNI, Dirección y Teléfono, y presioná 'Crear Cliente'.")

    add_sub_header(doc, "B. Cómo Elegir la Lista de Precios")
    add_bullet(doc, "Lista por Defecto:", "Al seleccionar al cliente, se carga su lista habitual.")
    add_bullet(doc, "Cambio de Lista:", "Podés cambiarla desde el selector desplegable (ej: Mayorista, Minorista, Especial).")
    add_bullet(doc, "Recálculo Automático:", "Si ya tenías productos en el carrito, todos los precios se ajustan instantáneamente a la nueva lista.")

    add_sub_header(doc, "C. Cómo Cargar Productos desde el Catálogo")
    add_step(doc, "1", "Abrir Catálogo:", "Tocá el botón 'CATÁLOGO'.")
    add_step(doc, "2", "Buscar / Filtrar:", "Buscá por nombre, código de artículo o filtrá por Marca y Categoría.")
    add_step(doc, "3", "Fotos en Zoom:", "Tocá la miniatura de la foto para ver la imagen ampliada en alta resolución.")
    add_step(doc, "4", "Ajustar Cantidades:", "Usá los botones '+' y '-'. La app controla el stock disponible en tiempo real.")
    add_step(doc, "5", "Cerrar Catálogo:", "Tocá el botón inferior 'CERRAR CATÁLOGO' para volver al pedido.")

    add_sub_header(doc, "D. Cómo Cargar Combos Promocionales")
    add_step(doc, "1", "Acceso a Promos:", "Tocá el botón 'COMBOS & PROMOS' o la pestaña 'Combos' en la barra inferior.")
    add_step(doc, "2", "Revisión:", "Revisá los artículos incluidos, el porcentaje de descuento y el precio total.")
    add_step(doc, "3", "Agregar:", "Tocá 'Agregar Combo al Carrito'. Los productos se desglosarán en el pedido con la etiqueta 'Combo: [Nombre]'.")

    add_sub_header(doc, "E. Cómo Aplicar Descuentos por Producto")
    add_bullet(doc, "Casilla Dto %:", "En cada ítem del carrito tenés un campo con el ícono '%'.")
    add_bullet(doc, "Ingreso de porcentaje:", "Escribí el descuento (ej: '5' para 5%). El sistema recalcula el precio final respetando las reglas de redondeo.")

    add_sub_header(doc, "F. Fecha de Entrega y Notas de Despacho")
    add_step(doc, "1", "Revisar Pedido:", "Tocá el botón 'REVISAR PEDIDO' para abrir la pantalla de remito.")
    add_step(doc, "2", "Fecha Programada:", "Tocá los botones rápidos 'Hoy', 'Mañana' o seleccioná una fecha en el calendario.")
    add_step(doc, "3", "Notas para Depósito:", "Escribí indicaciones de entrega (ej: 'Entregar de mañana', 'Tocar timbre local 4').")
    add_step(doc, "4", "Confirmar y Enviar:", "Tocá el botón verde 'CONFIRMAR Y ENVIAR'.")

    # 3. Repartos y Cobranzas
    add_section_header(doc, "3", "Módulos de Repartos y Cobranzas", "Operaciones de entrega en destino y cobro de recibos en calle.")

    add_sub_header(doc, "Gestión de Entregas (Pestaña 'Repartos')")
    add_bullet(doc, "Ver Mapa:", "Abre Google Maps con la ubicación exacta del cliente.")
    add_bullet(doc, "Llamar:", "Inicia una llamada telefónica directa al número de contacto del cliente.")
    add_bullet(doc, "Botón 'Entregado':", "Confirma que la mercadería fue recibida y cierra el pedido.")
    add_bullet(doc, "Botón 'No Entregado':", "Permite registrar el motivo de incidencia (ej: Local cerrado, Cliente ausente) para reprogramarlo.")

    add_sub_header(doc, "Cobranza de Cuentas Corrientes (Pestaña 'Cobranzas')")
    add_step(doc, "1", "Seleccionar deudor:", "Elegí el cliente de la lista de cuentas con saldo pendiente.")
    add_step(doc, "2", "Factura a cobrar:", "Tocá 'Cobrar' en el comprobante adeudado.")
    add_step(doc, "3", "Registrar pago:", "Ingresá el monto cobrado, elegí si fue Efectivo, Tarjeta o Transferencia y presioná 'CONFIRMAR PAGO'.")

    # 4. Guía de Botones
    add_section_header(doc, "4", "Referencia Completa de Botones de la App", "Descripción y función de cada elemento interactivo.")

    btn_table = doc.add_table(rows=1, cols=3)
    btn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    btn_table.autofit = False
    
    b_cells = btn_table.rows[0].cells
    b_cells[0].width = Inches(1.6)
    b_cells[1].width = Inches(1.8)
    b_cells[2].width = Inches(3.1)
    
    b_headers = ["Botón / Elemento", "Ubicación", "Función Principal"]
    for i, h in enumerate(b_headers):
        set_cell_background(b_cells[i], "1E1B4B")
        set_cell_margins(b_cells[i], top=100, bottom=100, left=120, right=120)
        p = b_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    btn_data = [
        ("Indicador WIFI / OFFLINE", "Barra Superior", "Muestra estado de red y avisa si se guardó localmente."),
        ("Refrescar (Flechas)", "Barra Superior", "Fuerza sincronización de pedidos, actualiza stocks y listas."),
        ("Salir (Rojo)", "Barra Superior", "Cierra la sesión del vendedor de manera segura."),
        ("➕ Nuevo", "Barra Inferior", "Abre la pantalla de confección y carga de pedidos."),
        ("✨ Combos", "Barra Inferior", "Muestra el catálogo de combos y promociones vigentes."),
        ("🚚 Repartos", "Barra Inferior", "Muestra las entregas asignadas al vendedor/repartidor."),
        ("🕒 Pedidos", "Barra Inferior", "Historial de pedidos emitidos y sus estados."),
        ("🔖 Cobranzas", "Barra Inferior", "Módulo de cobro de facturas y cuentas corrientes en la calle."),
        ("CATÁLOGO", "Pantalla Nuevo Pedido", "Abre el buscador visual de productos con stock y filtros."),
        ("REVISAR PEDIDO", "Pantalla Nuevo Pedido", "Abre el remito para fijar fecha de entrega, notas y confirmar."),
        ("Ver Mapa", "Pestaña Repartos", "Abre la ubicación en Google Maps para guiar al repartidor."),
        ("Llamar", "Pestaña Repartos", "Inicia una llamada telefónica directa al cliente."),
        ("✅ Entregado", "Pestaña Repartos", "Confirma la entrega exitosa del pedido."),
        ("❌ No Entregado", "Pestaña Repartos", "Registra motivo de no entrega (cliente ausente, local cerrado, etc.)."),
        ("🚚 Listo Entrega", "Historial de Pedidos", "Pasa el pedido a armado para repartirlo de inmediato."),
        ("Editar / Anular", "Historial de Pedidos", "Permite modificar o cancelar pedidos pendientes liberando el stock.")
    ]

    for row_idx, (bname, bloc, bfun) in enumerate(btn_data):
        row_cells = btn_table.add_row().cells
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([bname, bloc, bfun]):
            cell = row_cells[i]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            set_cell_border(cell, bottom=dict(sz=2, val="single", color="E2E8F0"))
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.name = "Segoe UI"
            r.font.size = Pt(8.5)
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif i == 1:
                r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            else:
                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()

    output_path = r"c:\xampp\htdocs\tomassidok\Manual_de_Uso_App_Vendedores.docx"
    doc.save(output_path)
    print(f"Manual de la App guardado en: {output_path}")

if __name__ == "__main__":
    generar_manual_sistema()
    generar_manual_app()
