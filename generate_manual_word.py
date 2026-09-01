import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def create_styled_manual():
    doc = Document()

    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Palette
    PRIMARY = RGBColor(37, 99, 235)      # #2563EB - Royal Blue
    DARK_NAVY = RGBColor(15, 23, 42)     # #0F172A
    SLATE = RGBColor(71, 85, 105)        # #475569
    EMERALD = RGBColor(16, 185, 129)     # #10B981
    MUTED_GRAY = RGBColor(100, 116, 139) # #64748B

    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def set_cell_border(cell, **kwargs):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", 4)}" w:space="{edge_data.get("space", 0)}" w:color="{edge_data.get("color", "auto")}"/>'
                tcBorders.append(parse_xml(tag))
            else:
                tag = f'<w:{edge} {nsdecls("w")} w:val="none"/>'
                tcBorders.append(parse_xml(tag))
        tcPr.append(tcBorders)

    def add_callout(text, title="NOTA IMPORTANTE", bg_color="EFF6FF", border_color="2563EB"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(6.8)
        
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
        set_cell_border(cell, 
                        left=dict(sz=24, val="single", color=border_color),
                        top=dict(sz=4, val="single", color="E2E8F0"),
                        bottom=dict(sz=4, val="single", color="E2E8F0"),
                        right=dict(sz=4, val="single", color="E2E8F0"))
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        run_title = p.add_run(f"📌 {title}: ")
        run_title.bold = True
        run_title.font.name = "Segoe UI"
        run_title.font.size = Pt(10.5)
        run_title.font.color.rgb = PRIMARY
        
        run_text = p.add_run(text)
        run_text.font.name = "Segoe UI"
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = DARK_NAVY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(13)
        run.font.color.rgb = DARK_NAVY
        return p

    def add_p(text, bold_prefix="", bullet=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bullet:
            p.paragraph_format.left_indent = Inches(0.25)
            run_b = p.add_run("• ")
            run_b.bold = True
            run_b.font.color.rgb = PRIMARY
        if bold_prefix:
            run_bp = p.add_run(bold_prefix + " ")
            run_bp.bold = True
            run_bp.font.name = "Segoe UI"
            run_bp.font.size = Pt(10.5)
            run_bp.font.color.rgb = DARK_NAVY
        run_t = p.add_run(text)
        run_t.font.name = "Segoe UI"
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = SLATE
        return p

    # --- COVER / TITLE BANNER ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    header_cell.width = Inches(6.8)
    set_cell_background(header_cell, "1E293B")
    set_cell_margins(header_cell, top=240, bottom=240, left=240, right=240)
    set_cell_border(header_cell)

    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun1 = hp.add_run("ONLYERP — MANUAL OPERATIVO Y CICLO DE VIDA\n")
    hrun1.bold = True
    hrun1.font.name = "Segoe UI"
    hrun1.font.size = Pt(18)
    hrun1.font.color.rgb = RGBColor(255, 255, 255)

    hrun2 = hp.add_run("Guía Estándar para Gestión de Pedidos, Hojas de Ruta y Cuentas Corrientes\n")
    hrun2.font.name = "Segoe UI"
    hrun2.font.size = Pt(12)
    hrun2.font.color.rgb = RGBColor(147, 197, 253)

    hrun3 = hp.add_run("NanoLabs SaaS Suite • Versión Oficial 1.0")
    hrun3.font.name = "Segoe UI"
    hrun3.font.size = Pt(9.5)
    hrun3.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- INTRODUCCIÓN ---
    add_heading_1("1. Introducción y Propósito del Manual")
    add_p("Este manual describe el flujo operativo estándar e integral de OnlyERP. Cualquier persona que se incorpore a la empresa (administradores, preventistas, personal de depósito o repartidores) debe seguir estos procedimientos para asegurar stock exacto, trazabilidad de entregas y cuadre financiero sin fisuras.")

    add_callout(
        "El sistema conecta automáticamente las 4 áreas neurálgicas del negocio: Ventas -> Depósito -> Reparto -> Caja/Finanzas. "
        "Ningún pedido se despacha sin control, y ninguna cobranza queda sin rendirse.",
        title="PRINCIPIO CLAVE DE OPERACIÓN",
        bg_color="F0FDF4",
        border_color="10B981"
    )

    # --- SECCIÓN 2: CICLO DE VIDA DEL PEDIDO ---
    add_heading_1("2. Ciclo de Vida de los Pedidos")
    add_heading_2("2.1 Canales de Ingreso de Pedidos")
    add_p("Preventistas toman pedidos en calle desde celular o tablet con listas de precios y bonificaciones asignadas.", bold_prefix="1. Preventa Móvil (/vendedor):", bullet=True)
    add_p("Clientes autorizados cargan sus pedidos directamente desde el catálogo digital mayorista.", bold_prefix="2. Portal B2B Mayorista (/portal-b2b):", bullet=True)
    add_p("Carga manual de pedidos o presupuestos por teléfono/mostrador en la administración.", bold_prefix="3. Mostrador y Administración (/pedidos):", bullet=True)

    add_heading_2("2.2 Estados del Pedido y Acciones Requeridas")
    
    # Table of states
    t_states = doc.add_table(rows=1, cols=4)
    t_states.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Estado", "Significado", "¿Quién actúa?", "Acción Obligatoria"]
    widths = [Inches(1.3), Inches(1.8), Inches(1.3), Inches(2.4)]
    
    for i, h in enumerate(headers):
        cell = t_states.cell(0, i)
        cell.width = widths[i]
        set_cell_background(cell, "2563EB")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    states_data = [
        ("PENDIENTE", "Recién ingresado.", "Administración", "Controlar stock, precios pactados y deuda crediticia del cliente."),
        ("APROBADO", "Validado comercialmente.", "Administración", "Al presionar 'Aprobar', el sistema reserva el stock y lo envía a depósito."),
        ("ARMADO", "Bultos embalados.", "Depósito", "Realiza el picking de mercadería y marca el pedido como 'Armado'."),
        ("EN_HOJA_DE_RUTA", "Asignado a chofer.", "Logística", "Se incluye en una Hoja de Ruta activa con remito y consolidado."),
        ("ENTREGADO", "Mercadería entregada.", "Chofer / Caja", "El chofer entrega el pedido y rinde el cobro (Efectivo/Cheque/CC)."),
        ("RECHAZADO", "Cliente no recibió.", "Chofer / Depósito", "La mercadería devuelta reingresa automáticamente al stock físico."),
        ("FACTURADO", "Circuito completado.", "Administración", "Emisión de Factura Electrónica AFIP o ticket definitivo.")
    ]

    for row_idx, (st, sig, who, act) in enumerate(states_data):
        row = t_states.add_row()
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate([st, sig, who, act]):
            c = row.cells[col_idx]
            c.width = widths[col_idx]
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=80, right=80)
            set_cell_border(c, top=dict(sz=2, val="single", color="CBD5E1"), bottom=dict(sz=2, val="single", color="CBD5E1"))
            cp = c.paragraphs[0]
            cr = cp.add_run(val)
            cr.font.name = "Segoe UI"
            cr.font.size = Pt(9)
            if col_idx == 0:
                cr.bold = True
                cr.font.color.rgb = PRIMARY
            else:
                cr.font.color.rgb = DARK_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECCIÓN 3: LOGÍSTICA Y HOJAS DE RUTA ---
    add_heading_1("3. Logística, Despacho y Hojas de Ruta")
    add_p("El módulo de logística conecta los pedidos armados en depósito con el chofer y la entrega en la puerta del cliente.")

    add_heading_2("3.1 Pasos para el Despacho de Reparto")
    add_p("Ingresar a /logistica/hojas-de-ruta -> 'Nueva Hoja de Ruta'. Seleccionar el Chofer/Repartidor y el vehículo correspondiente.", bold_prefix="Paso 1: Armado de Ruta:", bullet=True)
    add_p("El sistema lista todos los pedidos en estado ARMADO. El operador tilda los pedidos correspondientes a esa zona/recorrido.", bold_prefix="Paso 2: Selección de Pedidos:", bullet=True)
    add_p("En la hoja de ruta hacer clic en 'Imprimir Consolidado de Carga'. Este documento suma todos los artículos de todos los pedidos juntos (ej. 'Total a cargar: 50 fardos de gaseosa, 30 cajas de aceite'). Depósito carga el camión en 10 minutos sin leer pedido por pedido.", bold_prefix="Paso 3: Consolidado de Carga (Clave para Depósito):", bullet=True)
    add_p("Se imprime el itinerario para el chofer con direcciones, teléfonos, método de cobro pactado y los comprobantes/remitos.", bold_prefix="Paso 4: Hoja de Ruta para Chofer:", bullet=True)

    add_heading_2("3.2 Rendición y Cierre de Chofer en Caja Diaria")
    add_p("Al finalizar el recorrido, el chofer se presenta ante Administración/Caja con los valores recaudados y comprobantes firmados:")
    add_p("Se ingresa a /logistica/rendicion/[id]. Por cada pedido se confirma si fue Entregado Total, Entregado Parcial o Rechazado.", bullet=True)
    add_p("Se declara el desglose exacto: Efectivo, Cheques (Banco, N°, Vencimiento), Transferencias o Crédito a Cuenta Corriente.", bullet=True)
    add_p("Al confirmar la rendición, el dinero en efectivo entra AUTOMÁTICAMENTE a la Caja Diaria abierta (tipo 'COBRANZA_REPARTO'). Los cheques van a la Cartera de Valores y los ítems rechazados reingresan al stock sin necesidad de ajustes manuales.", bullet=True)

    add_callout(
        "NUNCA recibir dinero de un chofer sin realizar la rendición en el sistema. La rendición es el único documento que libera de responsabilidad al chofer e imputa el efectivo en la caja oficial.",
        title="CONTROL DE CAJA ESTRICTO",
        bg_color="FEF2F2",
        border_color="EF4444"
    )

    # --- SECCIÓN 4: CUENTAS CORRIENTES ---
    add_heading_1("4. Gestión de Cuentas Corrientes (Crédito a Clientes)")
    add_p("El módulo /cuentas-corrientes permite dar crédito a clientes habituales manteniendo control del riesgo de incobrabilidad.")

    add_heading_2("4.1 Parámetros de Crédito en la Ficha del Cliente (/clientes)")
    add_p("Monto máximo en pesos que el cliente puede adeudar (ej. $1.000.000). Si lo supera, el sistema bloquea ventas a cuenta corriente.", bold_prefix="• Límite de Crédito:", bullet=False)
    add_p("Plazo en días para cancelar facturas (ej. 7, 15 o 30 días). Pasada esa fecha, la factura figura como VENCIDA.", bold_prefix="• Días de Vencimiento:", bullet=False)
    add_p("Tope de bonificación que los vendedores pueden otorgarle sin requerir clave de supervisor.", bold_prefix="• Límite de Descuento:", bullet=False)

    add_heading_2("4.2 Registro de Cobranzas y Descuentos por Pronto Pago")
    add_p("1. Ir a /cuentas-corrientes, buscar al cliente y hacer clic en 'Ver Ficha'.")
    add_p("2. Presionar 'Registrar Pago / Abono'.")
    add_p("3. Ingresar el Monto Cobrado y Método de Pago (Efectivo, Transferencia, Cheque).")
    add_p("4. Descuento por Pronto Pago (Opcional): Si se acuerda un 5% de descuento por pago anticipado, el sistema cancela el 100% de la deuda pero ingresa a caja únicamente el efectivo neto recibido, dejando constancia en el recibo.")

    add_heading_2("4.3 Tratamiento de Cheques Rechazados")
    add_p("Si un banco rechaza un cheque recibido de un cliente:")
    add_p("1. Ir a /finanzas/cartera-valores, buscar el cheque y presionar 'Marcar Rechazado'.")
    add_p("2. El sistema automáticamente cambia el cheque a RECHAZADO y GENERA UN CARGO directo en la Cuenta Corriente del cliente por el valor del cheque, reabriendo la deuda de inmediato con número de cheque y motivo.")

    add_heading_2("4.4 Recálculo de Deuda por Inflación")
    add_p("Para facturas vencidas impagas durante meses con alta inflación:")
    add_p("En la ficha del cliente, presionar 'Recalcular por Inflación' en la factura vencida. El sistema recalcula los artículos a las listas de precios actuales del día, descontando los pagos previos que el cliente ya haya entregado.")

    # --- SECCIÓN 5: RESUMEN DE RESPONSABILIDADES ---
    add_heading_1("5. Matriz de Responsabilidades por Rol")
    
    t_roles = doc.add_table(rows=1, cols=3)
    t_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_headers = ["Rol / Puesto", "Responsabilidad Principal", "Pantallas Clave"]
    r_widths = [Inches(1.8), Inches(3.2), Inches(1.8)]

    for i, h in enumerate(r_headers):
        cell = t_roles.cell(0, i)
        cell.width = r_widths[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    roles_data = [
        ("Vendedor / Preventa", "Visitar clientes, relevar stock, cargar pedidos y cobrar facturas autorizadas.", "/vendedor\n/portal-b2b"),
        ("Administración / Caja", "Aprobar pedidos, facturar, gestionar cobranzas, abrir/cerrar caja y rendir choferes.", "/pedidos\n/caja\n/cuentas-corrientes"),
        ("Depósito / Picking", "Preparar bultos de pedidos aprobados, cargar camiones con el consolidado y controlar stock.", "/pedidos/armados\n/inventario"),
        ("Chofer / Reparto", "Entregar mercadería con hoja de ruta, cobrar según condición y rendir valores en caja.", "/logistica/hojas-de-ruta\n/logistica/rendicion")
    ]

    for row_idx, (rol, resp, scr) in enumerate(roles_data):
        row = t_roles.add_row()
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate([rol, resp, scr]):
            c = row.cells[col_idx]
            c.width = r_widths[col_idx]
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=80, right=80)
            set_cell_border(c, top=dict(sz=2, val="single", color="CBD5E1"), bottom=dict(sz=2, val="single", color="CBD5E1"))
            cp = c.paragraphs[0]
            cr = cp.add_run(val)
            cr.font.name = "Segoe UI"
            cr.font.size = Pt(9)
            if col_idx == 0:
                cr.bold = True
                cr.font.color.rgb = PRIMARY
            else:
                cr.font.color.rgb = DARK_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Save documents
    file_path_1 = "c:/xampp/htdocs/tomassidok/Manual_Operativo_Ciclo_de_Vida_OnlyERP.docx"
    file_path_2 = "c:/xampp/htdocs/onlyerp/Manual_Operativo_Ciclo_de_Vida_OnlyERP.docx"
    doc.save(file_path_1)
    doc.save(file_path_2)
    print(f"Archivos Word generados exitosamente en:\n{file_path_1}\n{file_path_2}")

if __name__ == "__main__":
    create_styled_manual()
